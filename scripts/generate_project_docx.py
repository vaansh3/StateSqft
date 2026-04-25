"""One-off generator: StateSqft project explanation as .docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "docs" / "StateSqft_Project_Explanation.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("StateSqft — Project Explanation", 0)
    p = doc.add_paragraph()
    p.add_run(
        "This document summarizes the repository layout, the purpose of each major source file, "
        "and the end-to-end workflow from user login through data-backed LLM responses."
    )

    doc.add_heading("1. Project overview", level=1)
    doc.add_paragraph(
        "The project has three parts: (1) a Python ML package under src/bia652p for loading the "
        "metro inventory CSV and running offline baselines; (2) a FastAPI backend that verifies "
        "Supabase JWTs, loads matching rows from the same CSV, optionally fits simple statistical "
        "signals, and calls OpenAI with structured JSON output; (3) a Vite + React frontend that "
        "authenticates with Supabase in the browser and calls the API for chat."
    )

    doc.add_heading("2. End-to-end workflow", level=1)
    doc.add_paragraph(
        "The following describes what happens when a signed-in user sends a chat message."
    )

    flow = doc.add_paragraph()
    flow.style = "List Number"
    flow.add_run(
        "Browser: User signs in via Supabase Auth (LoginPage). Session is held client-side."
    )
    doc.add_paragraph(
        "Browser: ChatPage sends POST /api/chat with the message and optional prior-turn history, "
        "with Authorization: Bearer <Supabase access token>.",
        style="List Number",
    )
    doc.add_paragraph(
        "Backend: require_supabase_user (deps.py) validates the JWT using SUPABASE_JWT_SECRET "
        "and/or JWKS from SUPABASE_URL.",
        style="List Number",
    )
    doc.add_paragraph(
        "Backend: retrieve_inventory_for_chat (inventory_data.py) parses the user message and "
        "conversation context, resolves date windows and metro areas against the CSV, and builds "
        "a text block of real inventory numbers plus chart payloads.",
        style="List Number",
    )
    doc.add_paragraph(
        "Backend: complete_chat (chat_openai.py) sends the user message, data context, and "
        "history to OpenAI; the model returns Markdown plus structured fields for charts and verdict UI.",
        style="List Number",
    )
    doc.add_paragraph(
        "Backend: Returns JSON (ChatResponseViz). If Supabase DB logging is enabled, the turn may "
        "be stored via supabase_service.",
        style="List Number",
    )
    doc.add_paragraph(
        "Browser: AssistantMessage renders Markdown, charts (Recharts), and structured tiles.",
        style="List Number",
    )

    doc.add_heading("2.1 Workflow diagram (text)", level=2)
    mono = doc.add_paragraph()
    mono.style = "Intense Quote"
    run = mono.add_run(
        "+----------+     JWT      +----------+     CSV + OpenAI    +----------+\n"
        "| Browser  | -----------> | FastAPI  | -------------------> | OpenAI   |\n"
        "| (React)  | <----------- | (Python) | <------------------- | API      |\n"
        "+----------+   JSON reply +----------+                        +----------+\n"
        "      ^                         |\n"
        "      |                         v\n"
        "      |                   +----------+\n"
        "      +---- Supabase -----| Supabase |\n"
        "           Auth           | (Auth;   \n"
        "                           | optional DB)\n"
        "                           +----------+\n"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("3. Root directory", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Path"
    hdr[1].text = "Role"
    for c in hdr:
        _set_cell_shading(c, "D9E2F3")
    rows_root = [
        ("README.md", "Human-readable setup: Supabase, env files, how to run backend + frontend."),
        ("pyproject.toml", "Root Python package statesqft: ML deps for src/bia652p."),
        ("requirements.txt", "Pinned-style deps for the ML package; use with pip install -e . from root."),
        (".gitignore", "Git ignore rules for the repo."),
        ("Metro_invt_fs_uc_sfrcondo_sm_month.csv", "Zillow-style monthly for-sale inventory by metro (default data file for API and baselines)."),
        ("statesqft-current-architecture.png", "Optional architecture diagram asset (if present)."),
    ]
    for path, role in rows_root:
        row = tbl.add_row().cells
        row[0].text = path
        row[1].text = role
    doc.add_paragraph()

    doc.add_heading("4. Python ML package — src/bia652p/", level=1)
    tbl2 = doc.add_table(rows=1, cols=2)
    h2 = tbl2.rows[0].cells
    h2[0].text = "File"
    h2[1].text = "Purpose"
    for c in h2:
        _set_cell_shading(c, "D9E2F3")
    ml_files = [
        ("__init__.py", "Package marker; may export version or public API."),
        ("config.py", "PROJECT_ROOT and default DATA_CSV path to the metro inventory file."),
        ("data/load.py", "Loads the metro CSV into tidy long format for modeling."),
        ("data/__init__.py", "Subpackage marker for data loaders."),
        ("features.py", "Feature engineering: seasonality (sin/cos month), lags, next-month targets, momentum flags."),
        ("evaluate.py", "Metrics helpers: regression (RMSE, MAE, R²) and classification (accuracy, F1, etc.)."),
    ]
    for f, purpose in ml_files:
        r = tbl2.add_row().cells
        r[0].text = f
        r[1].text = purpose
    doc.add_paragraph()

    doc.add_heading("5. Scripts — scripts/", level=1)
    tbl3 = doc.add_table(rows=1, cols=2)
    h3 = tbl3.rows[0].cells
    h3[0].text = "File"
    h3[1].text = "Purpose"
    for c in h3:
        _set_cell_shading(c, "D9E2F3")
    scr = [
        ("baseline.py", "Runs offline baseline models on the inventory CSV (see README)."),
        ("run_baselines.py", "Thin wrapper that executes baseline.py (backward-compatible entry point)."),
        ("supabase_chat_messages.sql", "Optional SQL for Supabase tables to persist chat history server-side."),
    ]
    for f, purpose in scr:
        r = tbl3.add_row().cells
        r[0].text = f
        r[1].text = purpose
    doc.add_paragraph()

    doc.add_heading("6. Backend — backend/app/", level=1)
    tbl4 = doc.add_table(rows=1, cols=2)
    h4 = tbl4.rows[0].cells
    h4[0].text = "File"
    h4[1].text = "Purpose"
    for c in h4:
        _set_cell_shading(c, "D9E2F3")
    be = [
        ("main.py", "FastAPI app: /health, /api/me, POST /api/chat. Wires JWT user, retrieval, OpenAI, response model."),
        ("settings.py", "Environment-driven settings: CSV path, Supabase URL/secrets, OpenAI key/model, CORS origins."),
        ("deps.py", "Dependency require_supabase_user: JWT verification (HS256 or JWKS)."),
        ("inventory_data.py", "CSV load/cache, metro and date parsing from natural language, retrieval for chat, chart specs, statsmodels-based signals when applicable."),
        ("chat_openai.py", "Async OpenAI client, system/user prompts, structured JSON + Markdown completion."),
        ("schemas_chat.py", "Pydantic models for API responses: structured blocks, charts, verdict payloads."),
        ("supabase_service.py", "Optional Supabase client (service role) for persisting chat turns when configured."),
        ("__init__.py", "App package marker."),
    ]
    for f, purpose in be:
        r = tbl4.add_row().cells
        r[0].text = f
        r[1].text = purpose
    doc.add_paragraph()
    doc.add_paragraph(
        "backend/pyproject.toml declares the FastAPI backend package and its dependencies (FastAPI, uvicorn, "
        "pandas, statsmodels, openai, supabase, etc.). backend/.env (not committed) holds secrets."
    )

    doc.add_heading("7. Frontend — frontend/", level=1)
    doc.add_paragraph(
        "frontend/package.json — npm scripts (dev, build, preview) and React/Vite dependencies.\n"
        "frontend/vite.config.ts — Vite bundler configuration.\n"
        "frontend/tsconfig.json, tsconfig.node.json — TypeScript compiler options.\n"
        "frontend/index.html — SPA entry HTML.\n"
        "frontend/.env.example — Template for VITE_SUPABASE_* and VITE_API_URL."
    )
    tbl5 = doc.add_table(rows=1, cols=2)
    h5 = tbl5.rows[0].cells
    h5[0].text = "Path under src/"
    h5[1].text = "Purpose"
    for c in h5:
        _set_cell_shading(c, "D9E2F3")
    fe = [
        ("main.tsx", "React root: mounts App and global styles."),
        ("App.tsx", "react-router routes: /login, protected / chat, redirects."),
        ("index.css", "Global CSS variables (theme), assistant markdown, login shell, chat progress bar."),
        ("pages/LoginPage.tsx", "Supabase email/password login UI."),
        ("pages/ChatPage.tsx", "Chat transcript, send message, POST /api/chat with history, loading bar while waiting."),
        ("components/AssistantMessage.tsx", "Renders assistant reply: Markdown, Recharts, structured summary, verdict bar."),
        ("lib/supabase.ts", "Supabase browser client from Vite env vars."),
        ("useSession.ts", "Hook: current Supabase session and loading state."),
        ("types/chat.ts", "TypeScript types aligned with API chat responses."),
        ("vite-env.d.ts", "Vite env typings for import.meta.env."),
    ]
    for f, purpose in fe:
        r = tbl5.add_row().cells
        r[0].text = f
        r[1].text = purpose
    doc.add_paragraph()

    doc.add_heading("8. How to run (summary)", level=1)
    for line in [
        "Terminal A — Backend: cd backend && source .venv/bin/activate && pip install -e . && uvicorn app.main:app --reload --host 127.0.0.1 --port 8000",
        "Terminal B — Frontend: cd frontend && npm install && npm run dev",
        "Configure backend/.env and frontend/.env.local per README.md before running.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("9. Data flow (chat)", level=1)
    doc.add_paragraph(
        "User text → (optional) history for context → JWT on API → inventory_data selects CSV rows "
        "and builds numeric context → OpenAI produces answer grounded in that context → frontend displays "
        "text and visualizations. The metric is monthly for-sale inventory (not permits or new listings)."
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run(
        "\nGenerated for documentation. Regenerate with: python scripts/generate_project_docx.py"
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
